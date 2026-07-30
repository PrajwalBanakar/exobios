import io

from docx import Document as DocxDocument

from app.ingestion.exceptions import DocumentParseError
from app.ingestion.parsers.base import DocumentParser, ParsedDocument, ParsedPage


class DocxParser(DocumentParser):
    """Extracts paragraph text from a .docx file.

    python-docx exposes no real page boundaries (pagination is a rendering
    concern the format doesn't store), so the whole document is returned as
    a single logical page.
    """

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        try:
            document = DocxDocument(io.BytesIO(content))
        except Exception as exc:
            raise DocumentParseError(filename=filename, reason=str(exc)) from exc

        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        if not text.strip():
            raise DocumentParseError(filename=filename, reason="DOCX contains no extractable text")

        return ParsedDocument(pages=[ParsedPage(page_number=1, text=text)])
